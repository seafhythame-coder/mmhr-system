#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import json
from pathlib import Path

# مكتبات معالجة الملفات
try:
    import PyPDF2
    import pytesseract
    from PIL import Image
    from docx import Document
    import re
except ImportError as e:
    print(f"❌ خطأ: المكتبة {e.name} غير مثبتة")
    print("⚠️ شغل: pip install -r requirements.txt")
    sys.exit(1)

class MMHRProcessor:
    """معالج المستندات الذكي"""
    
    def __init__(self, document_id):
        self.document_id = document_id
        self.supported_formats = ['.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png', '.tiff']
    
    # ✅ 1. فتح PDF (محمي أو عادي)
    def extract_pdf_text(self, pdf_path, password=None):
        """استخراج نصوص من ملف PDF"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # التحقق من التشفير
                if reader.is_encrypted:
                    if password:
                        reader.decrypt(password)
                    else:
                        # محاولة فتح بكلمة سر فارغة
                        reader.decrypt('')
                
                text = ""
                total_pages = len(reader.pages)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- صفحة {page_num + 1} من {total_pages} ---\n"
                            text += page_text
                    except Exception as e:
                        text += f"\n[⚠️ خطأ في استخراج الصفحة {page_num + 1}: {str(e)}]\n"
                
                if not text.strip():
                    return f"⚠️ لم يتمكن من استخراج نصوص من PDF (قد يكون ملف صور)"
                
                return text
        
        except Exception as e:
            return f"❌ خطأ في فتح PDF: {str(e)}"
    
    # ✅ 2. قراءة الصور (OCR)
    def extract_image_text(self, image_path):
        """استخراج نصوص من صور باستخدام OCR"""
        try:
            image = Image.open(image_path)
            
            # محاولة قراءة العربية والإنجليزية
            try:
                text = pytesseract.image_to_string(image, lang='ara+eng')
            except:
                # إذا فشلت، جرب بدون العربية
                text = pytesseract.image_to_string(image)
            
            if not text.strip():
                return "⚠️ لم يتمكن من استخراج نصوص من الصورة"
            
            return text
        
        except Exception as e:
            return f"❌ خطأ في معالجة الصورة: {str(e)}"
    
    # ✅ 3. فتح Word
    def extract_word_text(self, docx_path):
        """استخراج نصوص من ملف Word"""
        try:
            doc = Document(docx_path)
            
            text = ""
            
            # استخراج النصوص من الفقرات
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # استخراج النصوص من الجداول
            for table in doc.tables:
                text += "\n--- جدول ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
            
            if not text.strip():
                return "⚠️ لم يتمكن من استخراج نصوص من Word"
            
            return text
        
        except Exception as e:
            return f"❌ خطأ في فتح Word: {str(e)}"
    
    # ✅ 4. تنظيف النص العربي
    def clean_arabic_text(self, text):
        """تنظيف وتصحيح النصوص العربية"""
        
        # إزالة المسافات الزائدة
        text = re.sub(r'\s+', ' ', text)
        
        # إزالة الأحرف الغريبة (الاحتفاظ بالعربي والإنجليزي والأرقام والعلامات)
        text = re.sub(
            r'[^\u0600-\u06FF\u0621-\u064A\w\s\.\,\!\?\-\(\)\:\؛\n]', 
            '', 
            text
        )
        
        # توحيد الألف
        text = text.replace('أ', 'ا').replace('إ', 'ا').replace('آ', 'ا')
        
        # توحيد بعض الحروف
        text = text.replace('ؤ', 'و').replace('ئ', 'ي')
        
        return text.strip()
    
    # ✅ 5. فحص جودة النص
    def check_text_quality(self, text):
        """فحص جودة النص المستخرج"""
        
        quality_info = {
            'total_chars': len(text),
            'total_words': len(text.split()),
            'total_lines': text.count('\n') + 1,
            'has_arabic': bool(re.search(r'[\u0600-\u06FF]', text)),
            'has_english': bool(re.search(r'[a-zA-Z]', text)),
            'has_numbers': bool(re.search(r'\d', text)),
        }
        
        return quality_info
    
    # ✅ 6. معالجة شاملة
    def process(self, file_path):
        """معالجة شاملة للملف"""
        
        # التحقق من وجود الملف
        if not os.path.exists(file_path):
            return f"❌ الملف غير موجود: {file_path}"
        
        file_ext = Path(file_path).suffix.lower()
        file_size = os.path.getsize(file_path)
        
        print(f"\n📥 معالجة ملف جديد:")
        print(f"   📄 الاسم: {Path(file_path).name}")
        print(f"   📊 النوع: {file_ext}")
        print(f"   💾 الحجم: {file_size / 1024:.2f} KB")
        
        # ✅ استخراج النص بناءً على نوع الملف
        print(f"   🔄 جاري الاستخراج...")
        
        if file_ext == '.pdf':
            text = self.extract_pdf_text(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff']:
            text = self.extract_image_text(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self.extract_word_text(file_path)
        else:
            return f"❌ نوع ملف غير مدعوم: {file_ext}\n💡 الأنواع المدعومة: {', '.join(self.supported_formats)}"
        
        # التحقق من الأخطاء
        if text.startswith("❌"):
            return text
        
        # ✅ تنظيف النص
        print(f"   🧹 جاري التنظيف...")
        cleaned_text = self.clean_arabic_text(text)
        
        # ✅ فحص الجودة
        quality = self.check_text_quality(cleaned_text)
        
        # ✅ حفظ النتيجة
        output_path = str(Path(file_path).with_suffix('_processed.txt'))
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # كتابة معلومات المعالجة
                f.write("═" * 50 + "\n")
                f.write("📄 تقرير معالجة المستند\n")
                f.write("═" * 50 + "\n\n")
                
                f.write(f"📊 إحصائيات النص:\n")
                f.write(f"   • الأحرف: {quality['total_chars']}\n")
                f.write(f"   • الكلمات: {quality['total_words']}\n")
                f.write(f"   • الأسطر: {quality['total_lines']}\n")
                f.write(f"   • يحتوي عربي: {'✅ نعم' if quality['has_arabic'] else '❌ لا'}\n")
                f.write(f"   • يحتوي إنجليزي: {'✅ نعم' if quality['has_english'] else '❌ لا'}\n")
                f.write(f"   • يحتوي أرقام: {'✅ نعم' if quality['has_numbers'] else '❌ لا'}\n")
                
                f.write("\n" + "═" * 50 + "\n")
                f.write("📝 النص المعالج:\n")
                f.write("═" * 50 + "\n\n")
                
                f.write(cleaned_text)
        
        except Exception as e:
            return f"❌ خطأ في حفظ الملف: {str(e)}"
        
        print(f"   ✅ تم المعالجة بنجاح")
        print(f"   💾 الملف: {Path(output_path).name}\n")
        
        return cleaned_text

# ================================
# 🚀 نقطة الدخول الرئيسية
# ================================

if __name__ == "__main__":
    
    if len(sys.argv) < 2:
        print("❌ استخدام: python processor.py <file_path> [document_id]")
        print("💡 مثال: python processor.py document.pdf 123")
        sys.exit(1)
    
    file_path = sys.argv[1]
    document_id = sys.argv[2] if len(sys.argv) > 2 else "unknown"
    
    print("\n" + "="*50)
    print(f"🔧 معالج المستندات MMHR (ID: {document_id})")
    print("="*50)
    
    # إنشاء معالج جديد
    processor = MMHRProcessor(document_id)
    
    # معالجة الملف
    result = processor.process(file_path)
    
    # طباعة النتيجة (سيتم قراءتها من قبل Node.js)
    print(result)
    
    print("="*50 + "\n")
